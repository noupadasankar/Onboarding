"""DOCX document loader.

Uses python-docx to extract paragraph text and core document properties
from Office Open XML (.docx) files. Tables are included as tab-separated rows
so their content is not silently discarded.
"""
import io
import zipfile
from typing import Any

from docx import Document as DocxDocument
from docx.opc.exceptions import PackageNotFoundError

from app.loaders.base import BaseLoader
from app.loaders.exceptions import CorruptedFileError, EmptyFileError
from app.models.document import Document, FileType
from app.utils.metadata import clean_metadata


class DocxLoader(BaseLoader):
    @property
    def file_type(self) -> str:
        return "docx"

    def load(self, content: bytes, filename: str) -> Document:
        if not content:
            raise EmptyFileError(f"DOCX file '{filename}' is empty.")

        try:
            doc = DocxDocument(io.BytesIO(content))
        except (PackageNotFoundError, zipfile.BadZipFile) as exc:
            raise CorruptedFileError(
                f"Cannot read DOCX '{filename}': file is corrupted or not a valid .docx."
            ) from exc
        except Exception as exc:
            raise CorruptedFileError(
                f"Unexpected error reading DOCX '{filename}': {exc}"
            ) from exc

        # ── Extract text ──────────────────────────────────────────────────────
        parts: list[str] = []

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                parts.append(text)

        for table in doc.tables:
            for row in table.rows:
                row_text = "\t".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    parts.append(row_text)

        content_str = "\n".join(parts)

        # ── Extract metadata ──────────────────────────────────────────────────
        props = doc.core_properties
        headings = [
            p.text.strip()
            for p in doc.paragraphs
            if p.style.name.startswith("Heading") and p.text.strip()
        ]

        raw_meta: dict[str, Any] = {
            "title": props.title,
            "author": props.author,
            "created": str(props.created) if props.created else None,
            "modified": str(props.modified) if props.modified else None,
            "paragraph_count": len(doc.paragraphs),
            "word_count": len(content_str.split()),
            "headings": headings[:30],
        }

        return Document(
            filename=filename,
            file_type=FileType.DOCX,
            mime_type=(
                "application/vnd.openxmlformats-officedocument"
                ".wordprocessingml.document"
            ),
            content=content_str,
            metadata=clean_metadata(raw_meta),
            page_count=None,
            source=filename,
            size_bytes=len(content),
        )
