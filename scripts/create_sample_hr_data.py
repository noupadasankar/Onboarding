"""Generate sample binary HR documents for Increment 3.

Run once from the project root to create the PDF, DOCX, and XLSX sample
files that cannot be stored as plain text in the repository:

    python scripts/create_sample_hr_data.py

Prerequisites (already in pyproject.toml dev deps):
    pip install fpdf2 python-docx openpyxl

The text-based files (Onboarding_Process.txt, HR_FAQs.csv,
Holiday_Calendar.csv) are stored directly in data/raw/hr/ and do NOT
need this script.

Output: data/raw/hr/ (two directories up from scripts/)
"""
from __future__ import annotations

from pathlib import Path

# ── Paths ─────────────────────────────────────────────────────────────────────
PROJECT_ROOT = Path(__file__).resolve().parent.parent
HR_DIR = PROJECT_ROOT / "data" / "raw" / "hr"


def _ensure_dir() -> None:
    HR_DIR.mkdir(parents=True, exist_ok=True)
    print(f"Output directory: {HR_DIR}")


# ── PDF helpers ───────────────────────────────────────────────────────────────

def _write_pdf(dest: Path, sections: list[tuple[str, list[str]]]) -> None:
    """Write a multi-section PDF using fpdf2."""
    from fpdf import FPDF  # type: ignore[import-untyped]

    pdf = FPDF()
    for title, body in sections:
        pdf.add_page()
        pdf.set_font("Helvetica", style="B", size=16)
        pdf.cell(0, 12, title, new_x="LMARGIN", new_y="NEXT")
        pdf.set_font("Helvetica", size=11)
        pdf.ln(4)
        for para in body:
            pdf.multi_cell(0, 8, para)
            pdf.ln(3)
    pdf.output(str(dest))
    print(f"  Created: {dest.name} ({dest.stat().st_size:,} bytes)")


def make_employee_handbook() -> None:
    dest = HR_DIR / "Employee_Handbook.pdf"
    _write_pdf(dest, [
        ("EMPLOYEE HANDBOOK", [
            "Welcome to OptiAgent Consulting.",
            "This handbook outlines policies, procedures, and expectations for all employees.",
            "Please read it carefully and acknowledge receipt via the HR portal.",
        ]),
        ("CODE OF CONDUCT", [
            "All employees must act with integrity, respect, and professionalism at all times.",
            "Discrimination, harassment, and bullying are strictly prohibited.",
            "Report violations to your HR Business Partner immediately.",
            "Breaches of the Code of Conduct may result in disciplinary action up to and "
            "including dismissal.",
        ]),
        ("WORKING HOURS", [
            "Standard working hours are 09:00 to 17:30, Monday to Friday.",
            "Flexible working arrangements may be agreed with your line manager in writing.",
            "Overtime must be pre-approved by your line manager and is compensated at "
            "1.5x the standard hourly rate.",
        ]),
        ("LEAVE ENTITLEMENTS", [
            "Annual Leave: 20 days per year (pro-rated in your first year).",
            "Sick Leave: Statutory sick pay applies; a medical certificate is required after "
            "3 consecutive days of absence.",
            "Parental Leave: 16 weeks fully paid for the primary caregiver; "
            "4 weeks paid for the secondary caregiver.",
        ]),
        ("EXPENSES & TRAVEL", [
            "Submit all expense claims within 30 days of the expense date via the finance portal.",
            "Business travel must be pre-approved by your line manager.",
            "Economy class is standard for flights under 4 hours; business class requires "
            "director-level approval.",
        ]),
        ("PERFORMANCE REVIEWS", [
            "Performance reviews are conducted bi-annually in June and December.",
            "Objectives are set collaboratively at the start of each review cycle.",
            "Ratings: Exceeds Expectations | Meets Expectations | Needs Improvement.",
        ]),
        ("DISCIPLINARY PROCEDURE", [
            "Stage 1: Verbal warning (recorded in the personnel file).",
            "Stage 2: Written warning.",
            "Stage 3: Final written warning.",
            "Stage 4: Dismissal (with or without notice depending on the severity of the breach).",
            "Employees may be accompanied by a colleague or trade union representative at "
            "all disciplinary hearings.",
        ]),
        ("HEALTH & SAFETY", [
            "All employees must familiarise themselves with the emergency evacuation procedures.",
            "Report any health and safety concerns to the Facilities team at "
            "facilities@optiagent.com.",
            "DSE assessments are available for all office-based employees — contact HR to arrange.",
        ]),
    ])


def make_leave_policy() -> None:
    dest = HR_DIR / "Leave_Policy.pdf"
    _write_pdf(dest, [
        ("LEAVE POLICY", [
            "This policy governs all leave entitlements for permanent employees of "
            "OptiAgent Consulting.",
            "Effective Date: 1 January 2026 | Version: 4.1",
        ]),
        ("ANNUAL LEAVE", [
            "Entitlement: 20 working days per calendar year.",
            "Accrual: Leave accrues at 1.67 days per month in the first year.",
            "Carry-over: Up to 5 days may be carried into the following year. "
            "Carry-over must be used by 31 March or it is forfeited.",
            "Approval: Submit requests via the HR portal at least 5 working days in advance.",
            "Peak periods: No more than 3 team members may be on leave simultaneously "
            "without manager approval.",
        ]),
        ("SICK LEAVE", [
            "Self-certification: Up to 3 consecutive days without a medical certificate.",
            "Beyond 3 days: A GP certificate is required and must be submitted to HR.",
            "Return-to-work: Complete a return-to-work form on your first day back.",
            "Long-term absence: HR will engage an occupational health assessment after "
            "4 weeks of continuous absence.",
        ]),
        ("PARENTAL LEAVE", [
            "Primary caregiver: 16 weeks fully paid, followed by up to 26 weeks unpaid.",
            "Secondary caregiver: 4 weeks fully paid.",
            "Adoption leave mirrors the primary caregiver entitlement.",
            "Notification: Give at least 8 weeks written notice via the HR portal.",
            "KIT days: Up to 10 Keeping In Touch days are available during unpaid leave.",
        ]),
        ("COMPASSIONATE LEAVE", [
            "Up to 5 days fully paid leave for the death of an immediate family member.",
            "Up to 2 days for the death of a close friend or extended family member.",
            "Additional unpaid leave may be agreed with the line manager on a case-by-case basis.",
        ]),
        ("OTHER LEAVE TYPES", [
            "Study leave: Up to 5 days per year for approved job-related qualifications.",
            "Jury service: Full pay for the duration of jury service.",
            "Volunteer leave: 2 days per year for approved volunteering activities.",
            "Emergency dependant leave: Reasonable unpaid time off to deal with emergencies "
            "involving dependants.",
        ]),
    ])


# ── DOCX helper ───────────────────────────────────────────────────────────────

def make_company_policy() -> None:
    from docx import Document  # type: ignore[import-untyped]

    doc = Document()
    doc.core_properties.title = "Company Policy Manual"
    doc.core_properties.author = "HR Department"

    doc.add_heading("COMPANY POLICY MANUAL", 0)
    doc.add_paragraph("Effective Date: 1 January 2026 | Version: 3.2 | Owner: HR Department")
    doc.add_paragraph(
        "This manual sets out the policies that govern conduct, responsibilities, and "
        "entitlements for all employees of OptiAgent Consulting. Compliance is mandatory."
    )

    policies: list[tuple[str, list[str]]] = [
        ("1. EQUAL OPPORTUNITIES POLICY", [
            "OptiAgent Consulting is committed to providing equal opportunities in employment "
            "for all current and prospective employees, regardless of age, disability, gender "
            "reassignment, marriage or civil partnership, pregnancy or maternity, race, religion "
            "or belief, sex, or sexual orientation.",
            "All recruitment, promotion, and development decisions are made on the basis of "
            "merit and the requirements of the role.",
            "Complaints of discrimination should be raised with the HR Business Partner or "
            "anonymously via the ethics hotline.",
        ]),
        ("2. ANTI-HARASSMENT & BULLYING POLICY", [
            "The company has zero tolerance for harassment or bullying of any kind.",
            "Harassment includes any unwanted conduct related to a protected characteristic "
            "that violates a person's dignity or creates an intimidating, hostile, degrading, "
            "humiliating, or offensive environment.",
            "Bullying is defined as offensive, intimidating, malicious, or insulting behaviour "
            "that abuses power to undermine, humiliate, denigrate, or injure an individual.",
            "All complaints will be investigated confidentially, promptly, and impartially. "
            "Victimisation of complainants will itself be treated as a disciplinary matter.",
        ]),
        ("3. DATA PROTECTION POLICY", [
            "All employees must process personal data in accordance with the UK GDPR and the "
            "Data Protection Act 2018.",
            "Personal data must not be shared outside the company without appropriate legal "
            "basis and authorisation.",
            "Data breaches — whether actual or suspected — must be reported to the Data "
            "Protection Officer (DPO) within 24 hours of discovery.",
            "Employees who knowingly or recklessly misuse personal data face disciplinary "
            "action up to and including dismissal.",
        ]),
        ("4. SOCIAL MEDIA POLICY", [
            "Employees must not post confidential, commercially sensitive, or client-identifiable "
            "information on social media platforms.",
            "Personal opinions shared online must be clearly personal and must not be attributed "
            "to the company or made to appear as the company's position.",
            "Brand-related social media activity (e.g. LinkedIn posts on behalf of the company) "
            "must be approved in advance by the Marketing team.",
        ]),
        ("5. DISCIPLINARY PROCEDURE", [
            "The company's disciplinary procedure is designed to be fair, consistent, and "
            "transparent. The stages are:",
            "Stage 1: Verbal warning — recorded in the employee's personnel file.",
            "Stage 2: Written warning — remains active for 12 months.",
            "Stage 3: Final written warning — remains active for 24 months.",
            "Stage 4: Dismissal — with or without notice depending on the severity of the breach.",
            "Gross misconduct (e.g. theft, fraud, or serious breach of data protection) may "
            "result in summary dismissal without prior warning.",
        ]),
        ("6. EXPENSES POLICY", [
            "All business expenses must be pre-approved by a line manager before being incurred "
            "wherever possible.",
            "Expense claims must be submitted via the finance portal within 30 calendar days "
            "of the expense date. Late claims may be rejected.",
            "Original receipts (digital acceptable) are required for all items exceeding GBP 10.",
            "The following expense categories are covered: travel, accommodation, client "
            "entertainment, professional subscriptions, and training materials.",
        ]),
    ]

    for heading, paragraphs in policies:
        doc.add_heading(heading, 1)
        for para in paragraphs:
            doc.add_paragraph(para)
        doc.add_paragraph()

    dest = HR_DIR / "Company_Policy.docx"
    doc.save(str(dest))
    print(f"  Created: {dest.name} ({dest.stat().st_size:,} bytes)")


# ── XLSX helper ───────────────────────────────────────────────────────────────

def make_salary_grades() -> None:
    from openpyxl import Workbook  # type: ignore[import-untyped]

    wb = Workbook()

    # Sheet 1: Salary Grades
    ws1 = wb.active
    ws1.title = "Salary Grades"
    ws1.append(["Grade", "Band Title", "Min Salary (GBP)", "Mid Salary (GBP)", "Max Salary (GBP)"])
    for row in [
        ["A", "Junior Analyst",    28000,  33000,   38000],
        ["B", "Analyst",           38000,  44000,   50000],
        ["C", "Senior Analyst",    50000,  58000,   66000],
        ["D", "Manager",           66000,  76000,   86000],
        ["E", "Senior Manager",    86000,  98000,  110000],
        ["F", "Director",         110000, 128000,  146000],
        ["G", "Senior Director",  146000, 168000,  190000],
        ["H", "Partner",          190000, 225000,  260000],
    ]:
        ws1.append(row)

    # Sheet 2: Bonus Structure
    ws2 = wb.create_sheet("Bonus Structure")
    ws2.append(["Grade Range", "Performance Rating", "Bonus % of Base Salary"])
    for row in [
        ["A-C", "Exceeds Expectations",  15],
        ["A-C", "Meets Expectations",     8],
        ["A-C", "Needs Improvement",      0],
        ["D-E", "Exceeds Expectations",  20],
        ["D-E", "Meets Expectations",    12],
        ["D-E", "Needs Improvement",      0],
        ["F-G", "Exceeds Expectations",  30],
        ["F-G", "Meets Expectations",    18],
        ["F-G", "Needs Improvement",      0],
        ["H",   "Exceeds Expectations",  40],
        ["H",   "Meets Expectations",    25],
        ["H",   "Needs Improvement",      0],
    ]:
        ws2.append(row)

    # Sheet 3: Benefits by Grade
    ws3 = wb.create_sheet("Benefits by Grade")
    ws3.append(["Grade", "Pension (Company %)", "Holiday Days", "Life Insurance (x salary)", "Private Medical"])
    for row in [
        ["A-B", 5,  20, 2, "Basic"],
        ["C-D", 7,  22, 3, "Standard"],
        ["E-F", 8,  25, 4, "Enhanced"],
        ["G-H", 10, 30, 6, "Premium"],
    ]:
        ws3.append(row)

    dest = HR_DIR / "Salary_Grades.xlsx"
    wb.save(str(dest))
    print(f"  Created: {dest.name} ({dest.stat().st_size:,} bytes)")


# ── Main ──────────────────────────────────────────────────────────────────────

def main() -> None:
    print("Generating binary sample HR documents...")
    print(f"Project root: {PROJECT_ROOT}")
    _ensure_dir()

    errors: list[str] = []

    print("\nGenerating PDF files (requires fpdf2)...")
    for fn in [make_employee_handbook, make_leave_policy]:
        try:
            fn()
        except ImportError:
            msg = "  [SKIP] fpdf2 not installed. Run: pip install fpdf2"
            print(msg)
            errors.append(msg)
            break

    print("\nGenerating DOCX files (requires python-docx)...")
    try:
        make_company_policy()
    except ImportError:
        msg = "  [SKIP] python-docx not installed. Run: pip install python-docx"
        print(msg)
        errors.append(msg)

    print("\nGenerating XLSX files (requires openpyxl)...")
    try:
        make_salary_grades()
    except ImportError:
        msg = "  [SKIP] openpyxl not installed. Run: pip install openpyxl"
        print(msg)
        errors.append(msg)

    print(f"\nDone. Files written to: {HR_DIR}")
    if errors:
        print("\nSome files were skipped. Install missing packages and re-run:")
        print("  cd web/ai-service && pip install -e '.[dev]'")
    else:
        print("All binary sample files created successfully.")


if __name__ == "__main__":
    main()
